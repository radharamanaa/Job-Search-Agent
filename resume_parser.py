import io
import PyPDF2
import docx
import pandas as pd
import logging

# Configure logger for this module
logger = logging.getLogger(__name__)

def parse_resume(uploaded_file):
    """
    Parse uploaded resume file and extract text content.
    Supports PDF, DOCX, and TXT formats.

    Args:
        uploaded_file: The uploaded file object from Streamlit

    Returns:
        str: Extracted text from the resume
    """
    logger.info(f"Parsing resume file: {uploaded_file.name}")
    file_type = uploaded_file.name.split('.')[-1].lower()
    logger.debug(f"Detected file type: {file_type}")

    try:
        if file_type == 'pdf':
            logger.info("Parsing PDF file")
            return parse_pdf(uploaded_file)
        elif file_type == 'docx':
            logger.info("Parsing DOCX file")
            return parse_docx(uploaded_file)
        elif file_type == 'txt':
            logger.info("Parsing TXT file")
            return parse_txt(uploaded_file)
        else:
            logger.warning(f"Unsupported file format: {file_type}")
            return "Unsupported file format. Please upload a PDF, DOCX, or TXT file."
    except Exception as e:
        logger.error(f"Error parsing resume: {str(e)}", exc_info=True)
        return f"Error parsing resume: {str(e)}"

def parse_pdf(file):
    """Extract text from PDF file"""
    try:
        pdf_reader = PyPDF2.PdfReader(io.BytesIO(file.getvalue()))
        logger.debug(f"PDF has {len(pdf_reader.pages)} pages")

        text = ""
        for i, page in enumerate(pdf_reader.pages):
            page_text = page.extract_text()
            text += page_text + "\n"
            logger.debug(f"Extracted {len(page_text)} characters from page {i+1}")

        logger.info(f"Successfully extracted {len(text)} characters from PDF")
        return text
    except Exception as e:
        logger.error(f"Error parsing PDF: {str(e)}", exc_info=True)
        raise

def parse_docx(file):
    """Extract text from DOCX file"""
    try:
        doc = docx.Document(io.BytesIO(file.getvalue()))
        logger.debug(f"DOCX has {len(doc.paragraphs)} paragraphs")

        text = ""
        for para in doc.paragraphs:
            text += para.text + "\n"

        logger.info(f"Successfully extracted {len(text)} characters from DOCX")
        return text
    except Exception as e:
        logger.error(f"Error parsing DOCX: {str(e)}", exc_info=True)
        raise

def parse_txt(file):
    """Extract text from TXT file"""
    try:
        text = file.getvalue().decode("utf-8")
        logger.info(f"Successfully extracted {len(text)} characters from TXT file")
        return text
    except Exception as e:
        logger.error(f"Error parsing TXT: {str(e)}", exc_info=True)
        raise